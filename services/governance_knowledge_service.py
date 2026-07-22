"""
Governance Knowledge Service

Loads enterprise controls and ADRs from Word documents.

The Word documents remain the source of truth. Python parses the
document structure but does not hardcode control statements,
technology applicability or validation objectives.
"""

from pathlib import Path
import re

from docx import Document


class GovernanceKnowledgeService:

    CONTROL_HEADING_PATTERN = re.compile(
        r"^(?P<id>(?:CNTR|EC)-\d+)"
        r"\s+[–—-]\s+"
        r"(?P<title>.+)$"
    )

    def __init__(self):

        self.base_path = Path("governance")
        self.controls_path = (
            self.base_path / "controls"
        )
        self.adrs_path = (
            self.base_path / "adrs"
        )

    # -----------------------------------------------------
    # Public Method
    # -----------------------------------------------------

    def load_governance(self):

        return {
            "controls": self.load_controls(),
            "adrs": self.load_adrs()
        }

    # -----------------------------------------------------
    # Load Control Documents
    # -----------------------------------------------------

    def load_controls(self):

        control_libraries = []

        if not self.controls_path.exists():
            return control_libraries

        for file in sorted(
                self.controls_path.glob("*.docx")):

            if file.name.startswith("~$"):
                continue

            document = Document(file)

            controls = self._extract_controls(
                document
            )

            if controls:
                control_libraries.append({
                    "document": file.name,
                    "controls": controls
                })

        return control_libraries

    # -----------------------------------------------------
    # Extract Controls
    # -----------------------------------------------------

    def _extract_controls(
            self,
            document):
        """
        Extracts CNTR-* and EC-* controls.

        A paragraph is treated as a new control only when it
        matches a complete heading such as:

        CNTR-001 – Database Encryption
        EC-001 – Data at Rest Encryption

        Narrative text such as "EC-001 is intentionally
        technology-agnostic" is not treated as a new control.
        """

        controls = []

        current_control = None
        current_section = None

        for paragraph in document.paragraphs:

            raw_text = paragraph.text.strip()

            if not raw_text:
                continue

            lines = [
                line.strip()
                for line in raw_text.splitlines()
                if line.strip()
            ]

            for text in lines:

                heading_match = (
                    self.CONTROL_HEADING_PATTERN.match(
                        text
                    )
                )

                if heading_match:

                    if current_control:
                        controls.append(
                            current_control
                        )

                    control_id = (
                        heading_match
                        .group("id")
                        .strip()
                    )

                    control_name = (
                        heading_match
                        .group("title")
                        .strip()
                    )

                    current_control = {
                        "id": control_id,
                        "title": (
                            f"{control_id} – "
                            f"{control_name}"
                        ),
                        "name": control_name,
                        "controlType": (
                            "ENTERPRISE"
                            if control_id.startswith("EC-")
                            else "TECHNOLOGY"
                        ),
                        "statement": "",
                        "objective": "",
                        "implementationGuidance": ""
                    }

                    current_section = None
                    continue

                if current_control is None:
                    continue

                if text == "Control Statement":
                    current_section = "statement"
                    continue

                if text == "Control Objective":
                    current_section = "objective"
                    continue

                if text == "Implementation Guidance":
                    current_section = (
                        "implementationGuidance"
                    )
                    continue

                if text in {
                    "Enterprise Controls",
                    "Technology Controls",
                    "AI Interpretation Note"
                }:
                    current_section = None
                    continue

                if current_section:
                    current_control[
                        current_section
                    ] += text + " "

        if current_control:
            controls.append(
                current_control
            )

        for control in controls:
            for field in (
                "statement",
                "objective",
                "implementationGuidance"
            ):
                control[field] = (
                    control[field].strip()
                )

        return controls

    # -----------------------------------------------------
    # Load ADRs
    # -----------------------------------------------------

    def load_adrs(self):

        adrs = []

        if not self.adrs_path.exists():
            return adrs

        for file in sorted(
                self.adrs_path.glob("*.docx")):

            if file.name.startswith("~$"):
                continue

            adrs.append({
                "document": file.name,
                "content": self._read_document(
                    file
                )
            })

        return adrs

    # -----------------------------------------------------
    # Read Word Document
    # -----------------------------------------------------

    def _read_document(
            self,
            file_path):

        try:
            document = Document(file_path)

            paragraphs = []

            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if text:
                    paragraphs.append(text)

            return "\n".join(paragraphs)

        except Exception as ex:

            print(
                f"Unable to read "
                f"{file_path}: {ex}"
            )

            return ""
