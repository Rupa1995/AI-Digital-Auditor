from pathlib import Path
import re

from docx import Document


class GovernanceService:
    """
    Loads governance artefacts from the Governance Repository.

    Current MVP

        - Architecture Decision Records (ADRs)
        - Enterprise Technology Controls

    Future

        - Enterprise Standards
        - Audit Framework
        - Risk Library
    """

    def __init__(self):

        self.base_path = Path("governance")

    # ---------------------------------------------------------
    # Architecture Decision Records
    # ---------------------------------------------------------

    def load_adr(self, adr_name):

        file_path = (
            self.base_path /
            "adrs" /
            f"{adr_name}.docx"
        )

        if not file_path.exists():

            raise FileNotFoundError(
                f"{file_path} not found."
            )

        document = Document(file_path)

        content = []

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                content.append(paragraph.text)

        return {

            "id": adr_name,

            "documentType":
                "Architecture Decision Record",

            "title":
                content[2] if len(content) > 2 else adr_name,

            "version":
                "1.0",

            "content":
                "\n".join(content)

        }

    # ---------------------------------------------------------
    # Technology Controls
    # ---------------------------------------------------------

    def load_controls(self, control_document):

        """
        Load Enterprise Technology Controls from
        the Governance Repository.
        """

        file_path = (
            self.base_path /
            "controls" /
            "GCP" /
            f"{control_document}.docx"
        )

        if not file_path.exists():

            raise FileNotFoundError(
                f"{file_path} not found."
            )

        document = Document(file_path)

        controls = []

        control_pattern = re.compile(r"^DB-[A-Z]+-\d+")

        tables = document.tables

        table_index = 0

        for paragraph in document.paragraphs:

            heading = paragraph.text.strip()

            if not heading:

                continue

            #
            # Ignore title page and document headings
            #

            if not control_pattern.match(heading):

                continue

            #
            # Example:
            #
            # DB-AV-001 - Production Database High Availability
            #

            control_id, control_name = heading.split(
                " - ",
                1
            )

            if table_index >= len(tables):

                break

            table = tables[table_index]

            table_index += 1

            properties = {}

            for row in table.rows:

                key = row.cells[0].text.strip()

                value = row.cells[1].text.strip()

                properties[key] = value

            controls.append({

                "id": control_id,

                "name": control_name,

                "category":
                    properties.get("Category"),

                "requirement":
                    properties.get(
                        "Enterprise Requirement"
                    ),

                "expected":
                    properties.get(
                        "Expected Configuration"
                    ),

                "evidence":
                    properties.get(
                        "Evidence Source"
                    ),

                "risk":
                    properties.get(
                        "Risk Rating",
                        properties.get("Risk")
                    ),

                "businessImpact":
                    properties.get(
                        "Business Impact"
                    ),

                "recommendation":
                    properties.get(
                        "Recommendation"
                    )

            })

        return controls